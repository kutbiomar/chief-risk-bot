from __future__ import annotations

import io
import logging
import os
from types import SimpleNamespace
from typing import Any

import openpyxl
import pdfplumber
from docx import Document as DocxDocument

from backend.config import get_settings

MAX_PDF_PAGES = 250
MAX_DOCX_PARAGRAPHS = 400
MAX_XLSX_ROWS = 500
MAX_RAW_TEXT_BYTES = 60_000
MAX_EXTRACTED_ROWS = 200

logger = logging.getLogger(__name__)


def _normalize_text(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _truncate_text(text: str) -> tuple[str, bool]:
    raw = text.encode("utf-8")
    if len(raw) <= MAX_RAW_TEXT_BYTES:
        return text, False
    return raw[:MAX_RAW_TEXT_BYTES].decode("utf-8", errors="ignore"), True


def _extract_pdf(payload: bytes) -> tuple[str, list[list[str]], int, list[dict[str, Any]]]:
    lines: list[str] = []
    rows: list[list[str]] = []
    row_refs: list[dict[str, Any]] = []
    with pdfplumber.open(io.BytesIO(payload)) as pdf:
        if len(pdf.pages) > MAX_PDF_PAGES:
            raise ValueError("PDF exceeds the demo page-count limit")
        for page_index, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text:
                lines.append(page_text)
            for table_index, table in enumerate(page.extract_tables() or [], start=1):
                for row_index, row in enumerate(table, start=1):
                    cleaned = [_normalize_text(cell) or "" for cell in (row or [])]
                    if any(cleaned):
                        rows.append(cleaned)
                        row_refs.append(
                            {
                                "row_number": len(rows),
                                "page_number": page_index,
                                "table_index": table_index,
                                "table_row_index": row_index,
                                "cells": cleaned,
                            }
                        )
    return "\n".join(lines), rows[:MAX_EXTRACTED_ROWS], len(pdf.pages), row_refs[:MAX_EXTRACTED_ROWS]


def _is_test_runtime() -> bool:
    return bool(os.environ.get("PYTEST_CURRENT_TEST"))


def _extract_pdf_with_azure(payload: bytes) -> tuple[str, list[list[str]], int, list[dict[str, Any]]] | None:
    settings = get_settings()
    if _is_test_runtime():
        return None
    if not settings.azure_document_intelligence_endpoint or not settings.azure_document_intelligence_key:
        return None
    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential
    except Exception:
        logger.warning("Azure Document Intelligence SDK unavailable; using local PDF parser fallback")
        return None

    try:
        client = DocumentIntelligenceClient(
            endpoint=settings.azure_document_intelligence_endpoint,
            credential=AzureKeyCredential(settings.azure_document_intelligence_key),
        )
        poller = client.begin_analyze_document("prebuilt-layout", body=payload)
        result = poller.result()
    except Exception as exc:
        logger.warning("Azure Document Intelligence parse failed; using fallback: %s", exc)
        return None

    lines: list[str] = []
    rows: list[list[str]] = []
    row_refs: list[dict[str, Any]] = []
    for page in result.pages or []:
        for line in page.lines or []:
            content = _normalize_text(getattr(line, "content", None))
            if content:
                lines.append(content)
    for table_index, table in enumerate(result.tables or [], start=1):
        row_count = int(getattr(table, "row_count", 0) or 0)
        cells = getattr(table, "cells", None) or []
        matrix = [["" for _ in range(int(getattr(table, "column_count", 0) or 0))] for _ in range(row_count)]
        for cell in cells:
            row_index = int(getattr(cell, "row_index", 0) or 0)
            column_index = int(getattr(cell, "column_index", 0) or 0)
            if row_index < len(matrix) and column_index < len(matrix[row_index]):
                matrix[row_index][column_index] = _normalize_text(getattr(cell, "content", None)) or ""
        for row in matrix:
            if any(row):
                rows.append(row)
                row_refs.append(
                    {
                        "row_number": len(rows),
                        "page_number": int(getattr((getattr(table, "bounding_regions", None) or [SimpleNamespace(page_number=1)])[0], "page_number", 1)),
                        "table_index": table_index,
                        "table_row_index": len(row_refs) + 1,
                        "cells": row,
                    }
                )
    page_count = len(result.pages or [])
    return "\n".join(lines), rows[:MAX_EXTRACTED_ROWS], page_count, row_refs[:MAX_EXTRACTED_ROWS]


def _extract_docx(payload: bytes) -> tuple[str, list[list[str]], int, list[dict[str, Any]]]:
    doc = DocxDocument(io.BytesIO(payload))
    paragraphs = [_normalize_text(paragraph.text) or "" for paragraph in doc.paragraphs]
    if len(paragraphs) > MAX_DOCX_PARAGRAPHS:
        paragraphs = paragraphs[:MAX_DOCX_PARAGRAPHS]
    rows: list[list[str]] = []
    row_refs: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows[:MAX_EXTRACTED_ROWS], start=1):
            cleaned = [_normalize_text(cell.text) or "" for cell in row.cells]
            if any(cleaned):
                rows.append(cleaned)
                row_refs.append(
                    {
                        "row_number": len(rows),
                        "table_index": table_index,
                        "table_row_index": row_index,
                        "cells": cleaned,
                    }
                )
    return "\n".join(filter(None, paragraphs)), rows[:MAX_EXTRACTED_ROWS], len(doc.paragraphs), row_refs[:MAX_EXTRACTED_ROWS]


def _extract_xlsx(payload: bytes) -> tuple[str, list[list[str]], int, list[dict[str, Any]]]:
    workbook = openpyxl.load_workbook(io.BytesIO(payload), read_only=True, data_only=True)
    lines: list[str] = []
    rows: list[list[str]] = []
    row_refs: list[dict[str, Any]] = []
    sheet_count = 0
    for sheet in workbook.worksheets:
        sheet_count += 1
        lines.append(f"[Sheet] {sheet.title}")
        for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if index > MAX_XLSX_ROWS:
                break
            cleaned = [_normalize_text(cell) or "" for cell in row]
            if any(cleaned):
                rows.append(cleaned)
                lines.append(" | ".join(cleaned))
                row_refs.append(
                    {
                        "row_number": len(rows),
                        "sheet_name": sheet.title,
                        "sheet_row_index": index,
                        "cells": cleaned,
                    }
                )
    return "\n".join(lines), rows[:MAX_EXTRACTED_ROWS], sheet_count, row_refs[:MAX_EXTRACTED_ROWS]


def parse_document_layout(file_type: str, payload: bytes) -> dict[str, Any]:
    if file_type == "pdf":
        parsed = _extract_pdf_with_azure(payload)
        if parsed is None:
            raw_text, rows, page_count, row_refs = _extract_pdf(payload)
            parser_name = "azure-document-intelligence-fallback"
        else:
            raw_text, rows, page_count, row_refs = parsed
            parser_name = "azure-document-intelligence"
    elif file_type == "docx":
        raw_text, rows, page_count, row_refs = _extract_docx(payload)
        parser_name = "docx-layout-parser"
    elif file_type == "xlsx":
        raw_text, rows, page_count, row_refs = _extract_xlsx(payload)
        parser_name = "xlsx-layout-parser"
    else:
        raise ValueError("Unsupported document type")

    truncated_text, truncated = _truncate_text(raw_text)
    return {
        "raw_text": truncated_text,
        "rows": rows,
        "page_count": page_count,
        "row_refs": row_refs,
        "layout_artifact": {
            "parser": parser_name,
            "page_count": page_count,
            "row_count": len(rows),
            "row_refs": row_refs,
        },
        "raw_text_truncated": truncated,
        "parser": parser_name,
    }
