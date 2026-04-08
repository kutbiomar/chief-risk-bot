from __future__ import annotations

import hashlib
import io
import json
import re
from pathlib import Path
from uuid import uuid4

import openpyxl
import pdfplumber
from docx import Document as DocxDocument
from sqlalchemy import select, update
from sqlalchemy.orm import Session

from backend.models.content import Document, ExtractionResult
from backend.models.portfolio import PortfolioSnapshot, Position

DOCUMENT_LIMIT_BYTES = 50 * 1024 * 1024
MAX_PDF_PAGES = 250
MAX_DOCX_PARAGRAPHS = 400
MAX_XLSX_ROWS = 500
MAX_RAW_TEXT_BYTES = 60_000
MAX_EXTRACTED_ROWS = 200
DOCUMENT_TYPES = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
}
HEADER_ALIASES = {
    "ticker": "ticker",
    "symbol": "ticker",
    "security": "name",
    "security name": "name",
    "name": "name",
    "description": "name",
    "quantity": "quantity",
    "qty": "quantity",
    "shares": "quantity",
    "units": "quantity",
    "market value": "market_value_usd",
    "market value usd": "market_value_usd",
    "market_value_usd": "market_value_usd",
    "value": "market_value_usd",
    "asset class": "asset_class",
    "asset_class": "asset_class",
    "custodian": "custodian",
    "region": "geo_region",
    "geo region": "geo_region",
    "geo_region": "geo_region",
    "sector": "sector",
    "market segment": "market_segment",
    "market_segment": "market_segment",
    "currency": "position_currency",
    "position currency": "position_currency",
    "position_currency": "position_currency",
    "notes": "notes",
}
ASSET_CLASS_ALIASES = {
    "equity": "public_equity",
    "public equity": "public_equity",
    "public_equity": "public_equity",
    "stock": "public_equity",
    "fixed income": "fixed_income",
    "fixed_income": "fixed_income",
    "bond": "fixed_income",
    "private equity": "private_equity",
    "private_equity": "private_equity",
    "cash": "cash",
    "real estate": "real_estate",
    "real_estate": "real_estate",
    "alternative": "alternative",
    "alternatives": "alternative",
    "commodity": "commodity",
}
DEFAULT_EXTRACTED_ROW = {
    "ticker": None,
    "name": None,
    "quantity": None,
    "market_value_usd": None,
    "asset_class": None,
    "custodian": None,
    "geo_region": None,
    "sector": None,
    "market_segment": None,
    "position_currency": "USD",
    "notes": "Manual review required",
}

FRIENDLY_PARSE_ERRORS = {
    "pdf": "Unable to parse this PDF - it may be corrupted, password-protected, or not a valid PDF.",
    "docx": "Unable to parse this DOCX file - it may be corrupted, password-protected, or not a valid Word document.",
    "xlsx": "Unable to parse this XLSX file - it may be corrupted, password-protected, or not a valid Excel workbook.",
}


def _normalize_text(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _normalize_header(value: object) -> str | None:
    text = _normalize_text(value)
    if text is None:
        return None
    slug = re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()
    return HEADER_ALIASES.get(slug, HEADER_ALIASES.get(slug.replace(" ", "_")))


def _parse_float(value: object) -> float | None:
    text = _normalize_text(value)
    if text is None:
        return None
    cleaned = text.replace(",", "").replace("$", "").replace("(", "-").replace(")", "")
    cleaned = cleaned.replace("%", "")
    try:
        return float(cleaned)
    except ValueError:
        return None


def _normalize_asset_class(value: object) -> str | None:
    text = _normalize_text(value)
    if text is None:
        return None
    return ASSET_CLASS_ALIASES.get(text.lower(), text.lower().replace(" ", "_"))


def _looks_like_ticker(value: str | None) -> bool:
    if value is None:
        return False
    return bool(re.fullmatch(r"[A-Z0-9.\-]{1,12}", value.upper()))


def _truncate_text(text: str) -> tuple[str, bool]:
    raw = text.encode("utf-8")
    if len(raw) <= MAX_RAW_TEXT_BYTES:
        return text, False
    return raw[:MAX_RAW_TEXT_BYTES].decode("utf-8", errors="ignore"), True


def _ensure_storage_path(workspace_id: str, sha256: str, extension: str) -> Path:
    root = Path("backend/runtime/storage/documents") / workspace_id
    root.mkdir(parents=True, exist_ok=True)
    return root / f"{sha256[:16]}{extension}"


def _is_pdf_encrypted(payload: bytes) -> bool:
    return b"/Encrypt" in payload[:4096] or b"/Encrypt" in payload


def _validate_magic_bytes(extension: str, payload: bytes) -> None:
    if extension == ".pdf" and not payload.startswith(b"%PDF"):
        raise ValueError("PDF upload failed signature validation")
    if extension == ".docx" and not payload.startswith(b"PK"):
        raise ValueError("DOCX upload failed signature validation")
    if extension == ".xlsx" and not payload.startswith(b"PK"):
        raise ValueError("XLSX upload failed signature validation")


def validate_document_upload(filename: str, payload: bytes) -> tuple[str, str, str]:
    normalized = Path(filename).name
    extension = Path(normalized).suffix.lower()
    if extension not in DOCUMENT_TYPES:
        raise ValueError("Only PDF, DOCX, and XLSX files are accepted")
    if len(payload) > DOCUMENT_LIMIT_BYTES:
        raise ValueError("Document exceeds the 50MB demo limit")
    if ".." in normalized or "/" in normalized or "\\" in normalized:
        raise ValueError("Filename contains invalid traversal semantics")
    if extension == ".pdf" and _is_pdf_encrypted(payload):
        raise ValueError("Encrypted PDF files are not supported in demo mode")
    _validate_magic_bytes(extension, payload)
    return normalized, DOCUMENT_TYPES[extension], extension


def create_document(db: Session, *, workspace_id: str, uploaded_by: str, filename: str, payload: bytes, folder: str) -> Document:
    normalized, file_type, extension = validate_document_upload(filename, payload)
    digest = hashlib.sha256(payload).hexdigest()
    storage_path = _ensure_storage_path(workspace_id, digest, extension)
    storage_path.write_bytes(payload)
    document = Document(
        workspace_id=workspace_id,
        uploaded_by=uploaded_by,
        filename=normalized,
        file_type=file_type,
        file_size_bytes=len(payload),
        sha256=digest,
        storage_path=str(storage_path),
        folder=folder,
        malware_scan_status="clean",
        extraction_status="pending",
    )
    db.add(document)
    db.flush()
    return document


def _extract_pdf(payload: bytes) -> tuple[str, list[list[str]], int]:
    lines: list[str] = []
    rows: list[list[str]] = []
    with pdfplumber.open(io.BytesIO(payload)) as pdf:
        if len(pdf.pages) > MAX_PDF_PAGES:
            raise ValueError("PDF exceeds the demo page-count limit")
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                lines.append(page_text)
            for table in page.extract_tables() or []:
                for row in table:
                    cleaned = [_normalize_text(cell) or "" for cell in (row or [])]
                    if any(cleaned):
                        rows.append(cleaned)
    return "\n".join(lines), rows[:MAX_EXTRACTED_ROWS], len(pdf.pages)


def _extract_docx(payload: bytes) -> tuple[str, list[list[str]], int]:
    doc = DocxDocument(io.BytesIO(payload))
    paragraphs = [_normalize_text(paragraph.text) or "" for paragraph in doc.paragraphs]
    if len(paragraphs) > MAX_DOCX_PARAGRAPHS:
        paragraphs = paragraphs[:MAX_DOCX_PARAGRAPHS]
    rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows[:MAX_EXTRACTED_ROWS]:
            cleaned = [_normalize_text(cell.text) or "" for cell in row.cells]
            if any(cleaned):
                rows.append(cleaned)
    return "\n".join(filter(None, paragraphs)), rows[:MAX_EXTRACTED_ROWS], len(doc.paragraphs)


def _extract_xlsx(payload: bytes) -> tuple[str, list[list[str]], int]:
    workbook = openpyxl.load_workbook(io.BytesIO(payload), read_only=True, data_only=True)
    lines: list[str] = []
    rows: list[list[str]] = []
    sheet_count = 0
    for sheet in workbook.worksheets:
        sheet_count += 1
        lines.append(f"[Sheet] {sheet.title}")
        for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if index > MAX_XLSX_ROWS:
                break
            cleaned = [_normalize_text(cell) or "" for cell in row]
            if any(cleaned):
                rows.append(cleaned)
                lines.append(" | ".join(cleaned))
    return "\n".join(lines), rows[:MAX_EXTRACTED_ROWS], sheet_count


def _build_position(row: dict[str, object]) -> tuple[dict[str, object], dict[str, object]]:
    ticker = _normalize_text(row.get("ticker"))
    ticker = ticker.upper() if ticker else None
    name = _normalize_text(row.get("name"))
    quantity = _parse_float(row.get("quantity"))
    market_value = _parse_float(row.get("market_value_usd"))
    asset_class = _normalize_asset_class(row.get("asset_class"))
    position_currency = (_normalize_text(row.get("position_currency")) or "USD").upper()

    confidence = 0.2
    reasons: list[str] = []
    if ticker and _looks_like_ticker(ticker):
        confidence += 0.35
    elif ticker:
        reasons.append("ticker format uncertain")
    else:
        reasons.append("ticker missing")
    if quantity is not None:
        confidence += 0.2
    else:
        reasons.append("quantity missing")
    if market_value is not None:
        confidence += 0.15
    if asset_class:
        confidence += 0.1
    if name:
        confidence += 0.05

    notes = _normalize_text(row.get("notes"))
    if reasons:
        joined = "; ".join(reasons)
        notes = f"{notes}; {joined}" if notes else joined

    position = {
        "ticker": ticker,
        "name": name,
        "quantity": quantity,
        "market_value_usd": market_value,
        "asset_class": asset_class,
        "custodian": _normalize_text(row.get("custodian")),
        "geo_region": _normalize_text(row.get("geo_region")),
        "sector": _normalize_text(row.get("sector")),
        "market_segment": _normalize_text(row.get("market_segment")),
        "position_currency": position_currency,
        "notes": notes,
    }
    confidence_payload = {
        "row": 0,
        "confidence": round(min(confidence, 0.95), 2),
        "issues": reasons,
    }
    return position, confidence_payload


def _extract_positions_from_rows(rows: list[list[str]]) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    header_map: list[str | None] | None = None
    positions: list[dict[str, object]] = []
    confidence_rows: list[dict[str, object]] = []

    for raw in rows:
        if not any(cell.strip() for cell in raw):
            continue
        normalized_headers = [_normalize_header(cell) for cell in raw]
        if header_map is None and sum(1 for cell in normalized_headers if cell) >= 2:
            header_map = normalized_headers
            continue
        if header_map is None:
            continue

        record: dict[str, object] = {}
        for index, key in enumerate(header_map):
            if key is None or index >= len(raw):
                continue
            record[key] = raw[index]
        position, confidence = _build_position(record)
        if any(position[field] is not None for field in ("ticker", "name", "quantity", "market_value_usd")):
            confidence["row"] = len(positions) + 1
            positions.append(position)
            confidence_rows.append(confidence)
        if len(positions) >= MAX_EXTRACTED_ROWS:
            break

    return positions, confidence_rows


def parse_document(db: Session, document: Document) -> ExtractionResult:
    payload = Path(document.storage_path).read_bytes()

    if document.file_type == "pdf":
        try:
            raw_text, rows, page_count = _extract_pdf(payload)
        except Exception as exc:
            raise ValueError(FRIENDLY_PARSE_ERRORS["pdf"]) from exc
    elif document.file_type == "docx":
        try:
            raw_text, rows, page_count = _extract_docx(payload)
        except Exception as exc:
            raise ValueError(FRIENDLY_PARSE_ERRORS["docx"]) from exc
    elif document.file_type == "xlsx":
        try:
            raw_text, rows, page_count = _extract_xlsx(payload)
        except Exception as exc:
            raise ValueError(FRIENDLY_PARSE_ERRORS["xlsx"]) from exc
    else:
        raise ValueError("Unsupported document type")

    raw_text, truncated = _truncate_text(raw_text)
    positions, confidence = _extract_positions_from_rows(rows)

    if not positions:
        positions = [dict(DEFAULT_EXTRACTED_ROW, notes=f"Manual review required for {document.file_type} extraction")]
        confidence = [{"row": 1, "confidence": 0.25, "issues": ["no structured positions extracted"]}]

    needs_review = sum(
        1 for item in confidence if float(item.get("confidence", 0.0)) < 0.75
    )
    extraction = ExtractionResult(
        document_id=document.id,
        positions_json=json.dumps(positions, sort_keys=True),
        raw_text=raw_text or f"Parsed preview for {document.filename}",
        confidence_json=json.dumps(confidence, sort_keys=True),
        needs_review_count=needs_review,
        raw_text_truncated=truncated,
        extracted_row_count=len(positions),
        model=f"{document.file_type}-bounded-parser-v1",
        input_tokens=0,
        output_tokens=0,
    )
    db.add(extraction)
    db.flush()
    document.extraction_result_id = extraction.id
    document.extraction_status = "done"
    document.page_count = page_count
    db.flush()
    return extraction


def _clone_snapshot_from_positions(
    db: Session,
    *,
    workspace_id: str,
    uploaded_by: str,
    source: str,
    source_ref: str,
    positions: list[dict[str, object]],
) -> PortfolioSnapshot:
    current_snapshot = db.scalar(
        select(PortfolioSnapshot).where(
            PortfolioSnapshot.workspace_id == workspace_id,
            PortfolioSnapshot.is_current.is_(True),
        )
    )
    if current_snapshot is not None:
        db.execute(
            update(PortfolioSnapshot)
            .where(PortfolioSnapshot.id == current_snapshot.id)
            .values(is_current=False)
        )

    total_aum = round(
        sum(float(position.get("market_value_usd") or 0.0) for position in positions),
        2,
    )
    snapshot = PortfolioSnapshot(
        workspace_id=workspace_id,
        parent_snapshot_id=current_snapshot.id if current_snapshot is not None else None,
        uploaded_by=uploaded_by,
        source=source,
        source_ref=source_ref,
        position_count=len(positions),
        total_aum_usd=total_aum,
        is_current=True,
    )
    db.add(snapshot)
    db.flush()

    for position in positions:
        quantity = float(position.get("quantity") or 0.0)
        market_value = position.get("market_value_usd")
        market_value_usd = float(market_value) if market_value is not None else None
        db.add(
            Position(
                id=str(uuid4()),
                snapshot_id=snapshot.id,
                workspace_id=workspace_id,
                ticker=str(position.get("ticker") or "UNKNOWN").upper(),
                name=_normalize_text(position.get("name")),
                position_currency=str(position.get("position_currency") or "USD"),
                quantity=quantity,
                market_value_usd=market_value_usd,
                asset_class=str(position.get("asset_class") or "alternative"),
                geo_region=_normalize_text(position.get("geo_region")),
                sector=_normalize_text(position.get("sector")),
                market_segment=_normalize_text(position.get("market_segment")),
                custodian=_normalize_text(position.get("custodian")),
                price_source="document_approved",
                notes=_normalize_text(position.get("notes")),
            )
        )

    db.flush()
    return snapshot


def approve_document_extraction(db: Session, document: Document) -> PortfolioSnapshot:
    if document.extraction_result_id is None:
        raise ValueError("Document must be parsed before approval")
    extraction = db.get(ExtractionResult, document.extraction_result_id)
    if extraction is None:
        raise ValueError("Extraction result not found")

    raw_positions = json.loads(extraction.positions_json)
    approved_positions = [
        position
        for position in raw_positions
        if position.get("ticker") and position.get("quantity") is not None
    ]
    if not approved_positions:
        raise ValueError("No importable positions found in extraction result")

    snapshot = _clone_snapshot_from_positions(
        db,
        workspace_id=document.workspace_id,
        uploaded_by=document.uploaded_by,
        source="document_approved",
        source_ref=document.filename,
        positions=approved_positions,
    )
    document.tag = "reconciled"
    db.flush()
    return snapshot
